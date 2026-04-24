"""
build_docx.py — Word-Export gemäß Formatvorlage.md.

Erzeugt aus `Arbeitspapier_KI_Robotik_Besteuerung.md` ein Word-Dokument
im Corporate-Layout (A4, Navy/Stahlblau, Header/Footer, Akzentlinien
über Kapiteln, Sonderformat für die Deutschland-These in § 8.5).
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "Arbeitspapier_KI_Robotik_Besteuerung.md"
OUTPUT = ROOT / "Arbeitspapier_KI_Robotik_Besteuerung.docx"

PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT = RGBColor(0x2E, 0x75, 0xB6)
GREY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GREY = RGBColor(0x99, 0x99, 0x99)
H4_GREY = RGBColor(0x33, 0x33, 0x33)

DOC_TITLE = "Die Besteuerung von Künstlicher Intelligenz und Robotik als Ersatz menschlicher Arbeit"
DOC_SHORT_TITLE = "Arbeitspapier — Besteuerung von KI und Robotik"
DOC_ORG = "HIGL — Health Innovators Group Leipzig"
DOC_AUTHOR = "Björn Degenkolbe"
DOC_VERSION = "Version 9.0 — April 2026 — CC BY 4.0"

_INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
_INLINE_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_INLINE_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_runs(paragraph, text: str, base_font: str = "Arial", base_size: float = 9.5,
             color: RGBColor | None = None, italic: bool = False) -> None:
    """Render a markdown line as a sequence of runs (no nested formatting)."""

    text = _INLINE_LINK.sub(r"\1", text)  # links: keep label only
    tokens = re.split(r"(\*\*.+?\*\*|\*[^*]+\*)", text)
    for tok in tokens:
        if not tok:
            continue
        bold = False
        ital = italic
        content = tok
        if tok.startswith("**") and tok.endswith("**"):
            bold = True
            content = tok[2:-2]
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            ital = True
            content = tok[1:-1]
        run = paragraph.add_run(content)
        run.font.name = base_font
        run.font.size = Pt(base_size)
        run.bold = bold
        run.italic = ital
        if color is not None:
            run.font.color.rgb = color


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color_hex: str = "B0C4DE", size: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(size))
        e.set(qn("w:color"), color_hex)
        borders.append(e)
    tc_pr.append(borders)


def add_horizontal_rule(paragraph, color_hex: str, size_pt: float = 1.5) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(size_pt * 8)))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_paragraph_shading(paragraph, fill_hex: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    p_pr.append(shd)


def add_left_bar(paragraph, color_hex: str, size_pt: float = 3) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(int(size_pt * 8)))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color_hex)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def parse_table_block(lines: list[str], idx: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    while idx < len(lines) and lines[idx].lstrip().startswith("|"):
        raw = lines[idx].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        rows.append(cells)
        idx += 1
    if len(rows) >= 2 and all(set(c) <= set("-:") for c in rows[1]):
        rows.pop(1)
    return rows, idx


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = True
    for col_idx, header in enumerate(rows[0]):
        cell = table.rows[0].cells[col_idx]
        cell.text = ""
        para = cell.paragraphs[0]
        add_runs(para, header, base_size=7.5, color=RGBColor(255, 255, 255))
        for run in para.runs:
            run.bold = True
        set_cell_shading(cell, "1B3A5C")
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for r_idx, row in enumerate(rows[1:], start=1):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            add_runs(para, cell_text, base_size=7.5)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F7FB")


def add_quote(doc: Document, text: str, *, thesis: bool = False) -> None:
    para = doc.add_paragraph()
    add_paragraph_shading(para, "F0F5FA")
    bar_color = "1B3A5C" if thesis else "2E75B6"
    bar_size = 4 if thesis else 3
    add_left_bar(para, bar_color, size_pt=bar_size)
    para.paragraph_format.left_indent = Cm(0.8)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    add_runs(
        para,
        text,
        base_size=9.5 if thesis else 9,
        italic=True,
        color=PRIMARY if thesis else H4_GREY,
    )


def heading(doc: Document, text: str, level: int) -> None:
    para = doc.add_paragraph()
    if level == 2:
        add_runs(para, text, base_size=14, color=PRIMARY)
        for run in para.runs:
            run.bold = True
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(6)
    elif level == 3:
        add_runs(para, text, base_size=12, color=ACCENT)
        for run in para.runs:
            run.bold = True
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(4)
    elif level == 4:
        add_runs(para, text, base_size=10.5, color=H4_GREY)
        for run in para.runs:
            run.bold = True
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(2)


def chapter_accent(doc: Document, *, thick: bool = False) -> None:
    para = doc.add_paragraph()
    add_horizontal_rule(para, "2E75B6", size_pt=2 if thick else 1.5)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)


def configure_section(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.different_first_page_header_footer = True

    # Header (later pages)
    header = section.header
    h_para = header.paragraphs[0]
    h_para.text = ""
    tab_stops = h_para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(17.0), WD_ALIGN_PARAGRAPH.RIGHT)
    add_runs(h_para, f"{DOC_SHORT_TITLE}\t{DOC_ORG}", base_size=7.5, color=LIGHT_GREY)
    add_horizontal_rule(h_para, "CCCCCC", size_pt=0.5)

    # Footer (later pages)
    footer = section.footer
    f_para = footer.paragraphs[0]
    f_para.text = ""
    f_tab_stops = f_para.paragraph_format.tab_stops
    f_tab_stops.add_tab_stop(Cm(17.0), WD_ALIGN_PARAGRAPH.RIGHT)
    add_runs(f_para, f"{DOC_VERSION}\tSeite ", base_size=7.5, color=LIGHT_GREY)
    # Page-number field
    run = f_para.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "15")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "999999")
    r_pr.append(sz)
    r_pr.append(color)
    r.append(r_pr)
    fld.append(r)
    f_para._p.append(fld)


def build_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(p, "Arbeitspapier", base_size=16, color=PRIMARY)
    for run in p.runs:
        run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(p, DOC_TITLE, base_size=22, color=PRIMARY)
    for run in p.runs:
        run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(
        p,
        "Ökonomische, rechtliche und sozialpolitische Perspektiven",
        base_size=11,
        color=ACCENT,
    )

    for _ in range(8):
        doc.add_paragraph()

    for label in (DOC_AUTHOR, DOC_ORG, "Stand: April 2026", "Version 9.0",
                   "Lizenz: Creative Commons CC BY 4.0"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs(p, label, base_size=9, color=GREY)


def build(md_text: str) -> Document:
    doc = Document()
    configure_section(doc)

    # Default style adjustments
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9.5)
    style.paragraph_format.space_before = Pt(1)
    style.paragraph_format.space_after = Pt(2)

    build_cover(doc)
    doc.add_page_break()

    lines = md_text.splitlines()
    quote_buffer: list[str] = []
    list_buffer: list[str] = []
    paragraph_buffer: list[str] = []
    body_started = False
    in_thesis_quote = False
    skip_until_chapter1 = True

    def flush_paragraph() -> None:
        if paragraph_buffer:
            text = " ".join(paragraph_buffer).strip()
            if text:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                add_runs(p, text)
            paragraph_buffer.clear()

    def flush_list() -> None:
        if list_buffer:
            for item in list_buffer:
                p = doc.add_paragraph(style="List Bullet")
                add_runs(p, item)
            list_buffer.clear()

    def flush_quote() -> None:
        nonlocal in_thesis_quote
        if quote_buffer:
            text = " ".join(s.strip() for s in quote_buffer if s.strip())
            add_quote(doc, text, thesis=in_thesis_quote)
            quote_buffer.clear()
            in_thesis_quote = False

    def flush_all() -> None:
        flush_paragraph()
        flush_list()
        flush_quote()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if skip_until_chapter1:
            if re.match(r"^##\s+1\.\s", stripped):
                skip_until_chapter1 = False
            else:
                i += 1
                continue

        if stripped.startswith("## "):
            flush_all()
            heading_text = stripped[3:].strip()
            if body_started:
                doc.add_page_break()
            body_started = True
            chapter_accent(doc, thick=heading_text.lower().startswith("8."))
            heading(doc, heading_text, level=2)
            i += 1
            continue
        if stripped.startswith("### "):
            flush_all()
            heading(doc, stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("#### "):
            flush_all()
            heading(doc, stripped[5:].strip(), level=4)
            i += 1
            continue

        if stripped.startswith("|"):
            flush_all()
            rows, i = parse_table_block(lines, i)
            add_table(doc, rows)
            doc.add_paragraph()
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            flush_list()
            content = stripped[1:].strip()
            if content.startswith("**KI ist als eigenständiger dritter Produktionsfaktor"):
                in_thesis_quote = True
            quote_buffer.append(content)
            i += 1
            continue
        if quote_buffer and not stripped.startswith(">"):
            flush_quote()

        if stripped.startswith("- "):
            flush_paragraph()
            list_buffer.append(stripped[2:].strip())
            i += 1
            continue
        if list_buffer and not stripped.startswith("- "):
            flush_list()

        if not stripped or stripped == "---":
            flush_paragraph()
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_all()
    return doc


def main() -> int:
    md_text = SOURCE.read_text(encoding="utf-8")
    doc = build(md_text)
    doc.save(OUTPUT)
    print(f"DOCX geschrieben: {OUTPUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
